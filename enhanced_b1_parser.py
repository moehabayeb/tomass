#!/usr/bin/env python3
"""
Enhanced B1 module DOCX parser with comprehensive content extraction.
"""

import os
import json
import re
from docx import Document
from datetime import datetime

def extract_comprehensive_content(file_path):
    """Extract detailed content from DOCX with better parsing"""
    try:
        doc = Document(file_path)

        content = {
            'paragraphs': [],
            'tables': [],
            'full_text': "",
            'speaking_questions': [],
            'grammar_examples': [],
            'vocabulary_items': []
        }

        full_text_parts = []

        # Extract paragraphs with better encoding handling
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                # Clean up encoding issues
                text = text.encode('utf-8', errors='ignore').decode('utf-8')
                content['paragraphs'].append(text)
                full_text_parts.append(text)

        # Extract tables for grammar structures
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.strip().encode('utf-8', errors='ignore').decode('utf-8')
                    row_data.append(cell_text)
                if any(row_data):
                    table_data.append(row_data)
            if table_data:
                content['tables'].append(table_data)

        content['full_text'] = '\\n'.join(full_text_parts)

        # Extract speaking questions from content
        content['speaking_questions'] = extract_speaking_questions(content['paragraphs'])

        # Extract grammar examples
        content['grammar_examples'] = extract_grammar_examples(content['paragraphs'])

        # Extract vocabulary items
        content['vocabulary_items'] = extract_vocabulary_items(content['paragraphs'])

        return content

    except Exception as e:
        print(f"Error reading {file_path}: {e}")
        return {
            'paragraphs': [], 'tables': [], 'full_text': "",
            'speaking_questions': [], 'grammar_examples': [], 'vocabulary_items': []
        }

def extract_speaking_questions(paragraphs):
    """Extract actual speaking questions from content"""
    questions = []

    for para in paragraphs:
        # Look for question patterns
        if '?' in para and len(para) < 200:
            # Split on common delimiters to find individual questions
            potential_questions = re.split(r'[\\n\\r]+', para)

            for q in potential_questions:
                q = q.strip()
                if '?' in q and len(q) > 10 and len(q) < 150:
                    # Clean up the question
                    q = re.sub(r'^[\\d\\*\\-\\.\\s]+', '', q)  # Remove numbering
                    q = q.strip()
                    if q and q not in [item[0] for item in questions]:
                        # Generate appropriate answer
                        answer = generate_contextual_answer(q)
                        questions.append((q, answer))

    return questions[:40] if len(questions) >= 40 else questions

def extract_grammar_examples(paragraphs):
    """Extract example sentences for grammar"""
    examples = []

    for para in paragraphs:
        # Look for example indicators
        if any(keyword in para.lower() for keyword in ['example', 'örnek', 'sample', 'instance']):
            # Extract sentences that look like examples
            sentences = re.split(r'[.!?]+', para)
            for sent in sentences:
                sent = sent.strip()
                if 10 < len(sent) < 100 and not any(word in sent.lower() for word in ['example', 'örnek', 'students']):
                    examples.append(sent + '.')

    return examples[:10]

def extract_vocabulary_items(paragraphs):
    """Extract vocabulary words and phrases"""
    vocabulary = []

    for para in paragraphs:
        # Look for vocabulary sections
        if any(keyword in para.lower() for keyword in ['vocabulary', 'words', 'phrases', 'expressions']):
            # Extract word definitions or lists
            if ':' in para or '-' in para:
                items = re.split(r'[\\n\\r]+', para)
                for item in items:
                    item = item.strip()
                    if ':' in item or '-' in item and len(item) < 100:
                        vocabulary.append(item)

    return vocabulary[:20]

def generate_contextual_answer(question):
    """Generate contextually appropriate answers based on question type"""
    q_lower = question.lower()

    # Yes/No questions
    if question.startswith(('Are you', 'Do you', 'Can you', 'Have you', 'Will you', 'Would you')):
        return f"Yes, {question[4:].replace('?', '.').lower()}"
    elif question.startswith(('Is he', 'Is she', 'Is it')):
        return f"Yes, {question[3:].replace('?', '.').lower()}"
    elif question.startswith('Am I'):
        return f"Yes, you are{question[4:].replace('?', '.').lower()}"

    # Wh- questions
    elif question.startswith('What'):
        if 'do you' in q_lower:
            return "I usually do my best in that situation."
        elif 'is' in q_lower:
            return "It is something very important."
        else:
            return "That depends on the specific situation."

    elif question.startswith('How'):
        if 'long' in q_lower:
            return "For about three years now."
        elif 'often' in q_lower:
            return "I do it quite regularly."
        else:
            return "I do it very carefully and thoroughly."

    elif question.startswith('Where'):
        return "It is in a very convenient location."

    elif question.startswith('When'):
        return "It usually happens in the morning."

    elif question.startswith('Why'):
        return "Because it is very important for my goals."

    elif question.startswith('Who'):
        return "The person who helps me the most."

    # Default responses
    else:
        return "Yes, that is exactly right."

def generate_topic_specific_questions(module_num, title, content_data, vocabulary_items):
    """Generate 40 topic-specific questions based on module content"""

    # Use extracted questions first
    extracted_questions = content_data.get('speaking_questions', [])
    questions = []

    # Add extracted questions
    for q, a in extracted_questions:
        questions.append({"question": q, "answer": a})

    # Module-specific question generation based on topic
    topic = title.lower()

    if 101 <= module_num <= 110:  # Perfect tenses and modals
        base_templates = [
            ("Have you been practicing this grammar recently?", "Yes, I have been practicing it every day."),
            ("How long have you been studying this topic?", "I have been studying it for several weeks."),
            ("What had you learned before this lesson?", "I had learned the basic grammar rules."),
            ("Will you have mastered this by next month?", "Yes, I will have mastered it completely."),
            ("Can you explain this grammar rule?", "Yes, I can explain it clearly."),
            ("Must students practice this regularly?", "Yes, students must practice it every day."),
        ]
    elif 111 <= module_num <= 120:  # Conditionals and wish
        base_templates = [
            ("What would you do if you had more time?", "If I had more time, I would study more languages."),
            ("If you could change anything, what would it be?", "If I could change anything, I would improve my English."),
            ("Do you wish you had started learning earlier?", "Yes, I wish I had started learning earlier."),
            ("What would have happened if you hadn't studied?", "If I hadn't studied, I wouldn't have improved."),
            ("Would you mind if I asked you something?", "No, I wouldn't mind at all."),
            ("If only you could practice more, how would you feel?", "If only I could practice more, I would feel confident."),
        ]
    elif 121 <= module_num <= 130:  # Advanced grammar
        base_templates = [
            ("Are you used to speaking English daily?", "Yes, I am getting used to speaking it daily."),
            ("What do you need to have checked?", "I need to have my pronunciation checked."),
            ("Can you describe someone who inspired you?", "The person who inspired me was very dedicated."),
            ("Do you enjoy learning new expressions?", "Yes, I enjoy learning new expressions very much."),
            ("What makes you feel confident?", "Regular practice makes me feel confident."),
            ("How do you take care of your studies?", "I take care of my studies by practicing daily."),
        ]
    elif 131 <= module_num <= 140:  # Communication functions
        base_templates = [
            ("Could you tell me your opinion about this?", "In my opinion, this is very interesting."),
            ("What do you think about learning English?", "I think learning English is very important."),
            ("Do you agree with this approach?", "Yes, I completely agree with this approach."),
            ("How do you prefer to study?", "I prefer to study with interactive materials."),
            ("What happened first in your learning journey?", "First, I learned basic vocabulary and grammar."),
            ("Why do you think practice is important?", "Because practice helps us improve our skills."),
        ]
    else:  # 141-150: Vocabulary modules
        if 'work' in topic or 'job' in topic:
            base_templates = [
                ("What is your ideal job?", "My ideal job would be in education."),
                ("Do you prefer working in a team?", "Yes, I prefer working in a team environment."),
                ("What skills are important for your career?", "Communication skills are very important."),
                ("How do you handle workplace stress?", "I handle stress by staying organized."),
                ("What motivates you at work?", "Learning new things motivates me at work."),
            ]
        elif 'education' in topic or 'school' in topic:
            base_templates = [
                ("What subjects do you enjoy most?", "I enjoy language subjects the most."),
                ("How do you prepare for exams?", "I prepare by reviewing notes and practicing."),
                ("What makes a good teacher?", "A good teacher is patient and encouraging."),
                ("Do you prefer online or classroom learning?", "I prefer a combination of both methods."),
                ("What is your learning goal?", "My goal is to become fluent in English."),
            ]
        elif 'technology' in topic:
            base_templates = [
                ("How has technology changed education?", "Technology has made learning more interactive."),
                ("What apps do you use for learning?", "I use language learning apps regularly."),
                ("Do you think technology helps or distracts?", "I think technology helps when used properly."),
                ("What is your favorite educational website?", "My favorite site has interactive exercises."),
                ("How do you stay updated with technology?", "I read technology news regularly."),
            ]
        else:  # Other vocabulary topics
            base_templates = [
                (f"What do you know about {topic.split()[-1] if topic.split() else 'this topic'}?", f"I know quite a bit about {topic.split()[-1] if topic.split() else 'this topic'}."),
                (f"How often do you think about {topic.split()[-1] if topic.split() else 'this topic'}?", f"I think about {topic.split()[-1] if topic.split() else 'this topic'} quite often."),
                (f"What interests you most about {topic.split()[-1] if topic.split() else 'this topic'}?", f"What interests me most is the practical application."),
                ("Do you discuss these topics with friends?", "Yes, I often discuss these topics with friends."),
                ("How can this knowledge help you?", "This knowledge can help me in many situations."),
            ]

    # Fill remaining slots with base templates
    while len(questions) < 40:
        for template in base_templates:
            if len(questions) >= 40:
                break
            questions.append({"question": template[0], "answer": template[1]})

        # Add some variations if still need more
        if len(questions) < 40:
            questions.append({"question": "Can you give an example?", "answer": "Yes, I can give a clear example."})
            questions.append({"question": "How would you explain this to someone?", "answer": "I would explain it step by step."})
            questions.append({"question": "What is the most important point?", "answer": "The most important point is understanding the concept."})
            questions.append({"question": "Do you find this topic useful?", "answer": "Yes, I find this topic very useful."})

    return questions[:40]

def parse_enhanced_module_content(content_data, module_num, filename):
    """Parse content with comprehensive extraction"""
    paragraphs = content_data['paragraphs']
    tables = content_data['tables']
    vocabulary_items = content_data['vocabulary_items']

    # Extract title from filename
    title_match = re.search(r'Module_(\d+)_(.+)\.docx', filename)
    if title_match:
        topic = title_match.group(2).replace('_', ' ').title()
        title = f"Module {module_num} - {topic}"
    else:
        topic = "B1 Grammar"
        title = f"Module {module_num} - B1 Grammar"

    # Extract description based on module range
    if 101 <= module_num <= 120:
        description = f"Master {topic.lower()} - B1 intermediate grammar structure"
    elif 121 <= module_num <= 140:
        description = f"Learn {topic.lower()} - B1 advanced grammar patterns"
    else:  # 141-150
        description = f"Expand your {topic.lower()} - B1 vocabulary development"

    # Extract comprehensive intro
    intro_parts = []
    for para in paragraphs[:5]:
        if len(para) > 30 and any(word in para.lower() for word in ['objective', 'students', 'modül', 'learn']):
            intro_parts.append(para)

    intro = '\\n\\n'.join(intro_parts) if intro_parts else f"In this module, you will learn about {topic.lower()}. This is an important B1 level topic that will help you communicate more effectively in English."

    # Extract enhanced grammar tip
    tip = "Practice this topic regularly to improve your English proficiency."
    for para in paragraphs:
        if any(keyword in para.lower() for keyword in ['structure', 'form', 'rule', 'remember', 'tip']):
            tip = para[:200] + "..." if len(para) > 200 else para
            break

    # Process tables
    table_data = []
    if tables:
        table = tables[0]
        if len(table) > 1:
            headers = table[0] if table[0] else ["Structure", "Example"]
            for row in table[1:]:
                if len(row) >= 2:
                    table_data.append({
                        headers[0] if len(headers) > 0 else "Structure": row[0],
                        headers[1] if len(headers) > 1 else "Example": row[1] if len(row) > 1 else ""
                    })

    # Enhanced listening examples
    examples = content_data['grammar_examples'] if content_data['grammar_examples'] else [
        f"This is an example of {topic.lower()}.",
        f"Here's how we use {topic.lower()} in conversation.",
        f"Practice {topic.lower()} in daily situations."
    ]

    # Generate comprehensive speaking questions
    speaking_questions = generate_topic_specific_questions(module_num, title, content_data, vocabulary_items)

    return {
        "title": title,
        "description": description,
        "intro": intro,
        "tip": tip,
        "table": table_data,
        "listeningExamples": examples[:5],
        "speakingPractice": speaking_questions
    }

def main():
    """Enhanced processing for all 50 B1 modules"""
    extraction_dir = "temp_b1_extraction"
    modules_data = {}

    print("Processing ALL B1 modules 101-150 with enhanced parser...")

    # Process modules 101-150 (all 50 modules)
    for module_num in range(101, 151):
        # Find the actual file
        for filename in os.listdir(extraction_dir):
            if filename.startswith(f"Module_{module_num}_") and filename.endswith(".docx"):
                file_path = os.path.join(extraction_dir, filename)
                print(f"Processing {filename}...")

                # Extract comprehensive content
                content_data = extract_comprehensive_content(file_path)

                # Parse into structured data
                module_data = parse_enhanced_module_content(content_data, module_num, filename)
                modules_data[module_num] = module_data

                break
        else:
            print(f"Warning: Module {module_num} not found")

    # Create backup directory
    backup_dir = f"content_backups/{datetime.now().strftime('%Y%m%d_%H%M%S')}_complete_50_modules"
    os.makedirs(backup_dir, exist_ok=True)

    # Save comprehensive content backup
    with open(f"{backup_dir}/b1_all_50_modules.json", "w", encoding="utf-8") as f:
        json.dump(modules_data, f, indent=2, ensure_ascii=False)

    print(f"Processed {len(modules_data)} modules with enhanced parser")
    print(f"Backup saved to: {backup_dir}")

    return modules_data, backup_dir

if __name__ == "__main__":
    main()