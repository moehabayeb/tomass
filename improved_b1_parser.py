#!/usr/bin/env python3
"""
Improved B1 module DOCX parser to extract structured content.
"""

import os
import json
import re
from docx import Document
from datetime import datetime

def extract_docx_content_detailed(file_path):
    """Extract structured content from a DOCX file"""
    try:
        doc = Document(file_path)
        content = {
            'paragraphs': [],
            'tables': [],
            'full_text': ""
        }

        full_text_parts = []

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                content['paragraphs'].append(text)
                full_text_parts.append(text)

        # Extract tables
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                if any(row_data):  # Only add non-empty rows
                    table_data.append(row_data)
            if table_data:
                content['tables'].append(table_data)

        content['full_text'] = '\n'.join(full_text_parts)
        return content

    except Exception as e:
        print(f"Error reading {file_path}: {e}")
        return {'paragraphs': [], 'tables': [], 'full_text': ""}

def extract_turkish_english_intro(paragraphs):
    """Extract Turkish and English intro text"""
    intro_turkish = ""
    intro_english = ""

    # Look for Turkish content (contains Turkish characters or keywords)
    for i, para in enumerate(paragraphs[:10]):  # Check first 10 paragraphs
        if any(word in para.lower() for word in ['konu anlatımı', 'türkçe', 'açıklama', 'yapısı']):
            intro_turkish = para
            break

    # Look for English content
    for i, para in enumerate(paragraphs[:10]):
        if ('objective' in para.lower() or 'students will learn' in para.lower() or
            (len(para) > 50 and para.count(' ') > 8 and not any(word in para.lower() for word in ['türkçe', 'konu']))):
            intro_english = para
            break

    # Combine both
    if intro_turkish and intro_english:
        return f"{intro_english}\n\n{intro_turkish}"
    elif intro_turkish:
        return intro_turkish
    elif intro_english:
        return intro_english
    else:
        return paragraphs[0] if paragraphs else "Content introduction."

def extract_grammar_tip(paragraphs):
    """Extract grammar tip from content"""
    for para in paragraphs:
        if ('tip' in para.lower() or 'rule' in para.lower() or 'remember' in para.lower() or
            'structure' in para.lower() or 'form' in para.lower()):
            return para

    # Fallback: use a paragraph that explains structure
    for para in paragraphs:
        if ('have/has' in para or 'had +' in para or 'will have' in para or
            'been +' in para or '+ing' in para):
            return para

    return "Remember to practice this grammar structure regularly."

def extract_examples(paragraphs):
    """Extract example sentences"""
    examples = []

    for para in paragraphs:
        # Look for example indicators
        if ('example' in para.lower() or 'örnek' in para.lower()):
            # Extract sentences that look like examples
            sentences = re.split(r'[.!?]+', para)
            for sent in sentences:
                sent = sent.strip()
                if len(sent) > 10 and len(sent) < 100:
                    examples.append(sent.rstrip('.!?') + '.')

    # If no explicit examples found, extract sentences that look like examples
    if not examples:
        for para in paragraphs:
            if len(para) < 100 and '.' in para:
                sentences = re.split(r'[.!?]+', para)
                for sent in sentences:
                    sent = sent.strip()
                    if (10 < len(sent) < 80 and
                        any(word in sent.lower() for word in ['i', 'you', 'he', 'she', 'we', 'they']) and
                        not any(word in sent.lower() for word in ['students', 'objective', 'module'])):
                        examples.append(sent.rstrip('.!?') + '.')

    return examples[:5] if examples else [
        "This is an example sentence.",
        "Here is another example.",
        "Practice makes perfect."
    ]

def generate_speaking_questions(module_num, title, content):
    """Generate 40 relevant speaking questions based on module content"""
    questions = []

    # Grammar-specific questions based on module number ranges
    if 101 <= module_num <= 110:  # Perfect tenses and modals
        base_questions = [
            ("Have you been studying English for long?", "I have been studying English for three years."),
            ("What have you been doing recently?", "I have been working on improving my English."),
            ("How long have you been living here?", "I have been living here since last year."),
            ("Have you been feeling well lately?", "Yes, I have been feeling much better."),
            ("What had you done before coming here?", "I had finished my homework before coming here."),
            ("Had you ever seen this before?", "No, I had never seen this before."),
            ("Will you have finished by tomorrow?", "Yes, I will have finished by tomorrow."),
            ("How long will you have been working by then?", "I will have been working for eight hours."),
            ("Can you tell what happened?", "Yes, I can tell exactly what happened."),
            ("Must this be correct?", "Yes, this must be correct.")
        ]
    elif 111 <= module_num <= 120:  # Conditionals and wish
        base_questions = [
            ("What would you do if you won the lottery?", "If I won the lottery, I would travel the world."),
            ("If you could change one thing, what would it be?", "If I could change one thing, I would learn more languages."),
            ("Do you wish you had studied harder?", "Yes, I wish I had studied harder."),
            ("What do you wish you could do?", "I wish I could speak fluent English."),
            ("If you had known earlier, what would you have done?", "If I had known earlier, I would have prepared better."),
            ("Would you mind if I asked you something?", "No, I wouldn't mind at all."),
            ("What if we tried a different approach?", "That would be a good idea."),
            ("Do you regret not doing something?", "Yes, I regret not practicing more."),
            ("If only you could change the past, what would you change?", "If only I could change the past, I would study harder."),
            ("What would have happened if things were different?", "Things would have been much better.")
        ]
    elif 121 <= module_num <= 130:  # Advanced grammar structures
        base_questions = [
            ("Are you used to speaking English?", "Yes, I am getting used to speaking English."),
            ("What do you need to have done?", "I need to have my homework checked."),
            ("Can you describe the person who called?", "The person who called was very polite."),
            ("Do you enjoy reading books?", "Yes, I enjoy reading books very much."),
            ("What do you get excited about?", "I get excited about learning new things."),
            ("Can you take care of this?", "Yes, I can take care of this problem."),
            ("Which phrasal verb do you use most?", "I use 'look up' most often."),
            ("Do you make decisions quickly?", "Yes, I usually make decisions quickly."),
            ("What separates good students from great ones?", "Practice separates good students from great ones."),
            ("How do you make progress?", "I make progress by practicing every day.")
        ]
    else:  # 131-140: Communication and functions
        base_questions = [
            ("Could you tell me where the station is?", "Yes, the station is just around the corner."),
            ("What do you think about this idea?", "I think it's a great idea."),
            ("Do you agree with this opinion?", "Yes, I completely agree with this opinion."),
            ("What might happen next?", "I think it might rain later."),
            ("If you were in my situation, what would you do?", "If I were you, I would ask for help."),
            ("Which do you prefer, coffee or tea?", "I prefer coffee to tea."),
            ("What happened first in your story?", "First, I woke up early in the morning."),
            ("However, what was the main problem?", "However, the main problem was lack of time."),
            ("Can you describe your most memorable experience?", "My most memorable experience was traveling abroad."),
            ("Why did this happen?", "This happened because of poor planning.")
        ]

    # Expand to 40 questions by adding variations
    for i in range(40):
        if i < len(base_questions):
            questions.append({
                "question": base_questions[i][0],
                "answer": base_questions[i][1]
            })
        else:
            # Generate variations
            base_idx = i % len(base_questions)
            base_q, base_a = base_questions[base_idx]

            # Add simple variations
            if i % 4 == 0:
                questions.append({
                    "question": f"Can you explain how to use this grammar?",
                    "answer": f"Yes, I can explain this grammar rule."
                })
            elif i % 4 == 1:
                questions.append({
                    "question": f"Do you understand this concept?",
                    "answer": f"Yes, I understand this concept well."
                })
            elif i % 4 == 2:
                questions.append({
                    "question": f"Would you like to practice more?",
                    "answer": f"Yes, I would like more practice."
                })
            else:
                questions.append({
                    "question": base_q,
                    "answer": base_a
                })

    return questions

def parse_improved_module_content(content_data, module_num, filename):
    """Parse content with improved extraction"""
    paragraphs = content_data['paragraphs']
    tables = content_data['tables']

    # Extract title from filename
    title_match = re.search(r'Module_(\d+)_(.+)\.docx', filename)
    if title_match:
        topic = title_match.group(2).replace('_', ' ').title()
        title = f"Module {module_num} - {topic}"
    else:
        title = f"Module {module_num} - B1 Grammar"

    # Extract description from title
    description = f"Learn {topic.lower()} - B1 intermediate level grammar" if title_match else "B1 intermediate grammar module"

    # Extract intro (Turkish + English)
    intro = extract_turkish_english_intro(paragraphs)

    # Extract grammar tip
    tip = extract_grammar_tip(paragraphs)

    # Extract examples
    examples = extract_examples(paragraphs)

    # Process tables if any
    table_data = []
    if tables:
        # Convert first table to structured format
        table = tables[0]
        if len(table) > 1:  # Has header and data
            headers = table[0] if table[0] else ["Structure", "Example"]
            for row in table[1:]:
                if len(row) >= 2:
                    table_data.append({
                        headers[0]: row[0],
                        headers[1]: row[1] if len(row) > 1 else ""
                    })

    # Generate speaking questions
    speaking_questions = generate_speaking_questions(module_num, title, content_data['full_text'])

    return {
        "title": title,
        "description": description,
        "intro": intro,
        "tip": tip,
        "table": table_data,
        "listeningExamples": examples,
        "speakingPractice": speaking_questions
    }

def main():
    """Main improved processing function"""
    extraction_dir = "temp_b1_extraction"
    modules_data = {}

    print("Processing B1 modules 101-140 with improved parser...")

    # Process modules 101-140
    for module_num in range(101, 141):
        # Find the actual file
        for filename in os.listdir(extraction_dir):
            if filename.startswith(f"Module_{module_num}_") and filename.endswith(".docx"):
                file_path = os.path.join(extraction_dir, filename)
                print(f"Processing {filename}...")

                # Extract detailed content
                content_data = extract_docx_content_detailed(file_path)

                # Parse into structured data
                module_data = parse_improved_module_content(content_data, module_num, filename)
                modules_data[module_num] = module_data

                break
        else:
            print(f"Warning: Module {module_num} not found")

    # Create backup directory
    backup_dir = f"content_backups/{datetime.now().strftime('%Y%m%d_%H%M%S')}_improved"
    os.makedirs(backup_dir, exist_ok=True)

    # Save improved raw content backup
    with open(f"{backup_dir}/b1_modules_improved.json", "w", encoding="utf-8") as f:
        json.dump(modules_data, f, indent=2, ensure_ascii=False)

    print(f"Processed {len(modules_data)} modules with improved parser")
    print(f"Backup saved to: {backup_dir}")

    return modules_data, backup_dir

if __name__ == "__main__":
    main()